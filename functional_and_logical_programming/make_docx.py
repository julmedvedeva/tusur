#!/usr/bin/env python3
"""
Автосборка отчёта .docx из markdown по стандарту оформления ТУСУР.

Пайплайн:
  1. Подставляет полные листинги .py в Приложение А.
  2. pandoc: markdown -> docx (таблицы, заголовки, формулы -> уравнения Word).
  3. python-docx: применяет оформление ТУСУР (шрифт, поля, интервал, отступ,
     выравнивание, нумерация страниц).

Запуск:
  python3 make_docx.py report_lab2.md report_lab2.docx
"""
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Times New Roman"
SIZE = Pt(14)
CODE_SIZE = Pt(10)

def inject_listings(md: str, base: Path, stem: str) -> str:
    """Подставляет полный листинг исходника в Приложение А.

    Файл листинга выводится из имени отчёта:
      report_kr1.md -> kr1.lisp (Лисп, контрольная работа),
      report_lr1.md -> lr1.pl   (Пролог, лабораторная работа).
    """
    # Всё, что идёт после '## Приложение А', пересобираем заново.
    head, sep, _tail = md.partition("## Приложение А")
    if not sep:
        return md
    # report_kr1 -> kr1 ; report_lr1 -> lr1
    key = stem.replace("report_", "")
    num = key.replace("kr", "").replace("lr", "")
    if key.startswith("lr"):
        fname, lang, work = f"{key}.pl", "prolog", "Лабораторная работа"
    else:
        fname, lang, work = f"{key}.lisp", "lisp", "Контрольная работа"
    code = (base / fname).read_text(encoding="utf-8").rstrip()
    parts = [
        head,
        sep + " (обязательное) — Листинг программы\n",
        f"\n### А.1 {work} № {num} — все задачи ({fname})\n",
        f"\nФайл: `{fname}`\n",
        f"\n```{lang}\n{code}\n```\n",
    ]
    return "".join(parts)


def run_pandoc(md_path: Path, out_path: Path) -> None:
    subprocess.run(
        ["pandoc", str(md_path), "--from", "markdown", "--to", "docx",
         "-o", str(out_path)],
        check=True,
    )


def set_run_font(run, size=SIZE):
    run.font.name = FONT
    run.font.size = size
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), FONT)


def style_paragraph(p):
    name = (p.style.name or "").lower()
    pf = p.paragraph_format
    if name.startswith("heading") or name == "title":
        # Заголовки: TNR 14, полужирный, отступы сверху/снизу, без абз. отступа
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(12)
        pf.space_after = Pt(6)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        for r in p.runs:
            set_run_font(r)
            r.bold = True
    elif "source" in name or "verbatim" in name or "code" in name:
        # Листинги: моноширинный, мельче, одинарный интервал
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        for r in p.runs:
            r.font.size = CODE_SIZE
    else:
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.first_line_indent = Cm(1.25)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        for r in p.runs:
            set_run_font(r)


def add_table_borders(t):
    """Видимые границы всех ячеек через w:tblBorders."""
    tbl_pr = t._element.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")


def style_tables(doc):
    for t in doc.tables:
        add_table_borders(t)  # видимые границы всех ячеек
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.first_line_indent = Cm(0)
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    for r in p.runs:
                        set_run_font(r)


def set_margins(doc):
    for s in doc.sections:
        s.left_margin = Cm(3)
        s.right_margin = Cm(1)
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)


def add_page_numbers(doc):
    """Номер страницы снизу по центру."""
    for s in doc.sections:
        footer = s.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run = p.add_run()
        run._element.append(fld_begin)
        run._element.append(instr)
        run._element.append(fld_end)
        set_run_font(run)


def set_normal_style(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = SIZE
    rpr = st.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), FONT)


def main():
    src = Path(sys.argv[1] if len(sys.argv) > 1 else "report_lab2.md")
    out = Path(sys.argv[2] if len(sys.argv) > 2 else src.with_suffix(".docx"))
    base = src.resolve().parent

    md = src.read_text(encoding="utf-8")
    md = inject_listings(md, base, src.stem)
    tmp_md = base / "_report_full.md"
    tmp_md.write_text(md, encoding="utf-8")

    run_pandoc(tmp_md, out)

    doc = Document(str(out))
    set_normal_style(doc)
    set_margins(doc)
    for p in doc.paragraphs:
        style_paragraph(p)
    style_tables(doc)
    add_page_numbers(doc)
    doc.save(str(out))

    tmp_md.unlink(missing_ok=True)
    print(f"Готово: {out}")


if __name__ == "__main__":
    main()
