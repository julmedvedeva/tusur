#!/usr/bin/env python3
"""
Автосборка отчёта .docx из markdown по стандарту оформления ТУСУР
(ОС ТУСУР 01-2021 / СТО 02069326.1.01-2021).

Пайплайн:
  1. Подставляет полные листинги в Приложение А — берёт файлы, на которые
     в markdown уже стоят строки вида "Файл: `имя_файла`" под заголовками
     "### <ключ> <название>"; расширение файла определяет язык кодового блока.
  2. pandoc: markdown -> docx (таблицы, заголовки, формулы -> уравнения Word).
  3. python-docx: применяет оформление ТУСУР (шрифт, поля, интервал, отступ,
     выравнивание, заголовки по центру прописными, разрыв страницы перед
     каждым разделом, оформление титульного листа, нумерация страниц).

Запуск (из директории отчёта, рядом с report.md и листингами):
  python3 /path/to/utils/make_docx.py report_lab1.md report_lab1.docx

Титульный лист оформляется автоматически, если в markdown между строками-
метками `TITLEPAGE` и `END-TITLEPAGE` расположен текст титульного листа
(институт/тема — по центру, блок "Студент/Руководитель" — по правому краю;
см. report_lab1.md как пример). Строки этого блока, которые нужно прижать
вправо, перечислены в TITLEPAGE_RIGHT_BLOCK — отредактируйте под свой отчёт.

Важно: разделители-линии в markdown пишите как "***", а не "---" — "---"
после блока текста pandoc иногда трактует как открывающую границу YAML-
метаданных и падает с ошибкой парсинга YAML.
"""
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Times New Roman"
SIZE = Pt(14)
CODE_SIZE = Pt(10)

LISTING_LANG = {"py": "python", "lisp": "lisp", "pl": "prolog", "c": "c", "cpp": "cpp"}

LISTING_PATTERN = re.compile(r"### (\S+) (.+)\n\nФайл: `([^`]+)`", re.MULTILINE)

TITLEPAGE_RIGHT_BLOCK = {
    "Студент гр. ___________",
    "(подпись) Медведева Ю. Е.",
    "Руководитель",
    "к.т.н., доц. каф. АСУ",
    "(подпись) В. В. Романенко",
    "оценка ___________",
    "(дата)",
}


def inject_listings(md: str, base: Path) -> str:
    """Заменяет частичные листинги в Приложении А на полный код файлов."""
    head, sep, tail = md.partition("## ПРИЛОЖЕНИЕ А")
    if not sep:
        head, sep, tail = md.partition("## Приложение А")
    if not sep:
        return md
    nl = tail.find("\n")
    heading_rest, body = tail[:nl], tail[nl:]
    parts = [head, sep + heading_rest]
    for key, title, fname in LISTING_PATTERN.findall(body):
        path = base / Path(fname).name
        if not path.exists():
            path = base / fname
        code = path.read_text(encoding="utf-8").rstrip()
        lang = LISTING_LANG.get(path.suffix.lstrip("."), "")
        parts.append(f"\n### {key} {title}\n")
        parts.append(f"\nФайл: `{fname}`\n")
        parts.append(f"\n```{lang}\n{code}\n```\n")
    return "".join(parts)


def run_pandoc(md_path: Path, out_path: Path) -> None:
    subprocess.run(
        ["pandoc", str(md_path), "--from", "markdown", "--to", "docx",
         "-o", str(out_path)],
        check=True,
    )


def set_run_font(run, size=SIZE):
    run.font.name = FONT
    run.font.size = size
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), FONT)


def is_table_caption(text):
    return text.strip().startswith("Таблица ")


def style_paragraph(p):
    name = (p.style.name or "").lower()
    pf = p.paragraph_format
    text = p.text
    if name.startswith("heading") or name == "title":
        # Заголовки: ОС ТУСУР 01-2021 п.5.4.2 — по центру, полужирный, без абз. отступа
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(12)
        pf.space_after = Pt(6)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        # Раздел (Heading 2) начинается с нового листа (п.5.3.2); заголовок
        # титульного листа (Heading 1) новую страницу не открывает.
        if name == "heading 2":
            pf.page_break_before = True
        for r in p.runs:
            set_run_font(r)
            r.bold = True
    elif is_table_caption(text):
        # Название таблицы — над левым верхним углом, без абз. отступа (п.5.5.4)
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Cm(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        for r in p.runs:
            set_run_font(r)
    elif "source" in name or "verbatim" in name or "code" in name:
        # Листинги: моноширинный, мельче, одинарный интервал
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        for r in p.runs:
            r.font.size = CODE_SIZE
    else:
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.first_line_indent = Cm(1.25)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        for r in p.runs:
            set_run_font(r)


def add_table_borders(t):
    """Видимые границы всех ячеек через w:tblBorders."""
    tbl_pr = t._element.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")


def style_tables(doc):
    for t in doc.tables:
        add_table_borders(t)  # видимые границы всех ячеек
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.first_line_indent = Cm(0)
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    for r in p.runs:
                        set_run_font(r)


def set_margins(doc):
    for s in doc.sections:
        s.left_margin = Cm(3)
        s.right_margin = Cm(1.5)
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)


def add_page_numbers(doc):
    """Номер страницы снизу по центру; на титульном листе номер не ставится (п.5.11.1)."""
    for s in doc.sections:
        s.different_first_page_header_footer = True
        footer = s.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run = p.add_run()
        run._element.append(fld_begin)
        run._element.append(instr)
        run._element.append(fld_end)
        set_run_font(run)
        # Пустой колонтитул первой страницы (титульный лист)
        s.first_page_footer.is_linked_to_previous = False


def style_titlepage(doc):
    """Центрирует титульный лист по ОС ТУСУР 01-2021, прил. Б; удаляет служебные метки."""
    paras = doc.paragraphs
    start = end = None
    for i, p in enumerate(paras):
        if p.text.strip() == "TITLEPAGE":
            start = i
        elif p.text.strip() == "END-TITLEPAGE":
            end = i
            break
    if start is None or end is None:
        return
    for p in (paras[start], paras[end]):
        for r in list(p.runs):
            r.text = ""
    for p in paras[start + 1:end]:
        pf = p.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.alignment = (
            WD_ALIGN_PARAGRAPH.RIGHT if p.text.strip() in TITLEPAGE_RIGHT_BLOCK
            else WD_ALIGN_PARAGRAPH.CENTER
        )


def set_normal_style(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = SIZE
    rpr = st.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), FONT)


def main():
    if len(sys.argv) < 2:
        print("Usage: make_docx.py report.md [report.docx]", file=sys.stderr)
        sys.exit(1)
    src = Path(sys.argv[1])
    out = Path(sys.argv[2] if len(sys.argv) > 2 else src.with_suffix(".docx"))
    base = src.resolve().parent

    md = src.read_text(encoding="utf-8")
    md = inject_listings(md, base)
    tmp_md = base / "_report_full.md"
    tmp_md.write_text(md, encoding="utf-8")

    run_pandoc(tmp_md, out)

    doc = Document(str(out))
    set_normal_style(doc)
    set_margins(doc)
    for p in doc.paragraphs:
        style_paragraph(p)
    style_titlepage(doc)
    style_tables(doc)
    add_page_numbers(doc)
    doc.save(str(out))

    tmp_md.unlink(missing_ok=True)
    print(f"Готово: {out}")


if __name__ == "__main__":
    main()
